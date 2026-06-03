import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def parse_inline(paragraph, text, font_name="Times New Roman", font_size_pt=10, italic_all=False):
    # Matches bold (**text**), italic (*text*), math ($math$), and links ([text](url))
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|\$.*?\$|\[.*?\]\(.*?\))')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
            if italic_all:
                run.font.italic = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
        elif part.startswith('$') and part.endswith('$'):
            # Simple math formatting
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.name = 'Consolas'
            run.font.size = Pt(font_size_pt)
        elif part.startswith('[') and ']' in part and part.endswith(')'):
            match = re.search(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text = match.group(1)
                run = paragraph.add_run(link_text)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 238)
                run.font.name = font_name
                run.font.size = Pt(font_size_pt)
                if italic_all:
                    run.font.italic = True
        else:
            run = paragraph.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
            if italic_all:
                run.font.italic = True

def add_paragraph_ieee(doc, text, space_after_pt=4, line_spacing=1.0, first_line_indent_in=0.15, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent_in > 0:
        p.paragraph_format.first_line_indent = Inches(first_line_indent_in)
    
    parse_inline(p, text)
    return p

def main():
    workspace_dir = r"c:\Users\AYUSH K BHAT\OneDrive\Desktop\v2v-v2i-project-phase2"
    
    md_path = os.path.join(workspace_dir, "syntrix_v2x_ieee_paper.md")
    output_path = os.path.join(workspace_dir, "syntrix_v2x_ieee_paper.docx")

    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = docx.Document()

    # Define IEEE Page Setup for the first section
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # We will process line by line
    lines = md_content.split('\n')
    
    # Title & Metadata Section (1 column)
    i = 0
    title_raw = ""
    abstract_raw = ""
    keywords_raw = ""
    authors_raw = "Vishal, [Student 2 Name], [Student 3 Name], [Student 4 Name], [Student 5 Name]\nDepartment of Computer Science and Engineering, RV College of Engineering, Bengaluru, India"
    
    # First, let's find the Title, Abstract, and Keywords
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('# '):
            title_raw = line.lstrip('#').strip()
        elif line.startswith('## Abstract'):
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('**Keywords'):
                abstract_raw += lines[i].strip() + " "
                i += 1
            if i < len(lines) and lines[i].strip().startswith('**Keywords'):
                # Extract keywords
                kw_match = re.search(r'\*\*Keywords\u2014(.*?)\*\*', lines[i].strip())
                if not kw_match:
                    kw_match = re.search(r'\*\*Keywords\-(.*?)\*\*', lines[i].strip())
                if kw_match:
                    keywords_raw = kw_match.group(1).strip()
                else:
                    keywords_raw = lines[i].replace('**Keywords—', '').replace('**', '').strip()
            break
        i += 1

    # Add Title to Section 1
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run(title_raw)
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0, 0, 0)

    # Add Authors
    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_auth.paragraph_format.space_after = Pt(18)
    run_auth = p_auth.add_run(authors_raw)
    run_auth.font.name = 'Times New Roman'
    run_auth.font.size = Pt(10)
    run_auth.font.color.rgb = RGBColor(0, 0, 0)

    # Add Abstract
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs.paragraph_format.left_indent = Inches(0.25)
    p_abs.paragraph_format.right_indent = Inches(0.25)
    p_abs.paragraph_format.space_after = Pt(6)
    
    run_abs_tag = p_abs.add_run("Abstract—")
    run_abs_tag.bold = True
    run_abs_tag.font.italic = True
    run_abs_tag.font.name = 'Times New Roman'
    run_abs_tag.font.size = Pt(9)
    
    parse_inline(p_abs, abstract_raw, font_name="Times New Roman", font_size_pt=9, italic_all=True)

    # Add Keywords
    if keywords_raw:
        p_key = doc.add_paragraph()
        p_key.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_key.paragraph_format.left_indent = Inches(0.25)
        p_key.paragraph_format.right_indent = Inches(0.25)
        p_key.paragraph_format.space_after = Pt(18)
        
        run_key_tag = p_key.add_run("Keywords—")
        run_key_tag.bold = True
        run_key_tag.font.italic = True
        run_key_tag.font.name = 'Times New Roman'
        run_key_tag.font.size = Pt(9)
        
        run_key_val = p_key.add_run(keywords_raw)
        run_key_val.font.name = 'Times New Roman'
        run_key_val.font.size = Pt(9)

    # Add continuous section break for 2-column body
    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    body_section.top_margin = Inches(0.75)
    body_section.bottom_margin = Inches(0.75)
    body_section.left_margin = Inches(0.75)
    body_section.right_margin = Inches(0.75)
    
    # Configure 2 columns
    sectPr = body_section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '360') # 0.25 inches spacing between columns
    sectPr.append(cols)

    # Process remaining lines (Skip Title/Abstract/Keywords which we already added)
    i = 0
    in_abstract_block = False
    in_code_block = False
    code_text = ""
    table_lines = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip Title, Abstract, Keywords lines we processed manually
        if line.startswith('# ') or line.startswith('## Abstract') or line.startswith('**Keywords'):
            i += 1
            continue
        if abstract_raw and line in abstract_raw:
            i += 1
            continue
            
        # Handle Code Blocks
        if line.startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.1)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(code_text.strip())
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                in_code_block = False
                code_text = ""
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_text += lines[i] + "\n"
            i += 1
            continue

        # Handle Tables
        if line.startswith('|'):
            table_lines.append(lines[i])
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            if table_lines:
                parsed_rows = []
                for t_line in table_lines:
                    content = t_line.strip()
                    if content.startswith('|'):
                        content = content[1:]
                    if content.endswith('|'):
                        content = content[:-1]
                    cells = [c.strip() for c in content.split('|')]
                    is_sep = all(re.match(r'^[\s:-]*$', c) for c in cells)
                    if not is_sep:
                        parsed_rows.append(cells)
                
                if parsed_rows:
                    num_rows = len(parsed_rows)
                    num_cols = max(len(row) for row in parsed_rows)
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.style = 'Table Grid'
                    
                    for r_idx, row_data in enumerate(parsed_rows):
                        for c_idx in range(num_cols):
                            val = row_data[c_idx] if c_idx < len(row_data) else ""
                            val = val.replace('<br>', '\n')
                            
                            cell = table.cell(r_idx, c_idx)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.line_spacing = 1.0
                            
                            if r_idx == 0:
                                run = p.add_run(val)
                                run.bold = True
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(8.5)
                                set_cell_background(cell, "F2F2F2")
                            else:
                                parse_inline(p, val, font_size_pt=8.5)
                    doc.add_paragraph().paragraph_format.space_after = Pt(4)
                table_lines = []
            continue

        # Handle Headings
        if line.startswith('##'):
            # Level 1 (Section Heading like II. LITERATURE REVIEW)
            # Level 2 is '##' in the paper since title was '#'
            heading_text = line.lstrip('#').strip()
            
            # Check if Level 1 or Level 2 in IEEE context
            # We map "## I. Introduction" -> Level 1 (centered, bold, 10pt)
            # We map "### A. Heading" -> Level 2 (italic, left, 10pt)
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            if re.match(r'^[IXV]+\.\s', heading_text) or heading_text.startswith('References'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(heading_text.upper())
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(heading_text)
                run.font.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue
            
        if line.startswith('###'):
            # Level 2 (Sub-section Heading like A. Heading)
            heading_text = line.lstrip('#').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            
            run = p.add_run(heading_text)
            run.font.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # Handle Blockquotes (Equations or Figures)
        if line.startswith('>'):
            quote_text = line.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(4)
            
            if "Figure" in quote_text or "INSERT" in quote_text:
                # Figure caption style
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(quote_text.replace('**', ''))
                run.font.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(8.5)
            else:
                # Text blockquote
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                parse_inline(p, quote_text, font_size_pt=9)
                for run in p.runs:
                    run.font.italic = True
            i += 1
            continue

        # Handle List Bullets
        if line.startswith('- ') or line.startswith('* '):
            list_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run("•  ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            parse_inline(p, list_text, font_size_pt=10)
            i += 1
            continue

        # Handle Horizontal Rules
        if line == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Handle Equations $$ eq $$
        if line.startswith('$$') and line.endswith('$$'):
            eq_text = line[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(eq_text)
            run.font.name = 'Times New Roman'
            run.font.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Handle References or Normal Paragraphs
        if line:
            # Check if inside reference list (e.g. "1. Qin, X...")
            # References should be 8pt
            is_ref = re.match(r'^\d+\.\s', line)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if is_ref:
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                parse_inline(p, line, font_size_pt=8.0)
            else:
                p.paragraph_format.first_line_indent = Inches(0.15)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.0
                parse_inline(p, line, font_size_pt=10.0)
        else:
            pass
            
        i += 1

    try:
        doc.save(output_path)
        print(f"Successfully generated IEEE paper DOCX at: {output_path}")
    except PermissionError:
        alternative_path = os.path.join(workspace_dir, "syntrix_v2x_ieee_paper_final.docx")
        doc.save(alternative_path)
        print(f"Successfully generated alternative: {alternative_path}")

if __name__ == "__main__":
    main()
