import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# XML helper to set cell background color
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def parse_inline(paragraph, text):
    # Matches bold (**text**), italic (*text*), and links ([text](url))
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('[') and ']' in part and part.endswith(')'):
            match = re.search(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text = match.group(1)
                run = paragraph.add_run(link_text)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 238)
        else:
            paragraph.add_run(part)

def set_font_and_spacing(paragraph, font_name="Times New Roman", size_pt=12, line_spacing=1.15, space_after_pt=6):
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)

def add_paragraph_with_font(doc, text, style='Normal', space_after_pt=6):
    p = doc.add_paragraph(style=style)
    parse_inline(p, text)
    set_font_and_spacing(p, space_after_pt=space_after_pt)
    return p

def main():
    workspace_dir = r"c:\Users\AYUSH K BHAT\OneDrive\Desktop\v2v-v2i-project-phase2"
    template_path = os.path.join(workspace_dir, "sem2el.docx")
    output_path = os.path.join(workspace_dir, "v2x_project_report.docx")
    md_path = os.path.join(workspace_dir, "v2x_project_report.md")

    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist.")
        return

    # Load markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Load document
    if os.path.exists(template_path):
        print("Using sem2el.docx as template...")
        doc = docx.Document(template_path)
        # Clear body paragraphs
        p_elems = [p._element for p in doc.paragraphs]
        for p in p_elems:
            p.getparent().remove(p)
        # Clear existing tables (if any)
        t_elems = [t._element for t in doc.tables]
        for t in t_elems:
            t.getparent().remove(t)
    else:
        print("Template sem2el.docx not found. Creating a blank document.")
        doc = docx.Document()

    # Split markdown content into sections/blocks
    # We will process line by line
    lines = md_content.split('\n')
    
    i = 0
    in_code_block = False
    code_text = ""
    table_lines = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 1. Handle Code Blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(code_text.strip())
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                in_code_block = False
                code_text = ""
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_text += lines[i] + "\n"
            i += 1
            continue

        # 2. Handle Tables
        if line.startswith('|'):
            table_lines.append(lines[i])
            i += 1
            # Check if next line is also part of the table
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            # Process the collected table lines
            if table_lines:
                # Filter out separator row
                parsed_rows = []
                for t_line in table_lines:
                    # Clean outer pipes
                    content = t_line.strip()
                    if content.startswith('|'):
                        content = content[1:]
                    if content.endswith('|'):
                        content = content[:-1]
                    
                    # Split cells by pipe
                    cells = [c.strip() for c in content.split('|')]
                    
                    # Check if separator row (contains only dashes, colons, spaces)
                    is_sep = all(re.match(r'^[\s:-]*$', c) for c in cells)
                    if not is_sep:
                        parsed_rows.append(cells)
                
                if parsed_rows:
                    num_rows = len(parsed_rows)
                    num_cols = max(len(row) for row in parsed_rows)
                    
                    # Add table to docx
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.style = 'Table Grid'
                    
                    for r_idx, row_data in enumerate(parsed_rows):
                        for c_idx in range(num_cols):
                            if c_idx < len(row_data):
                                val = row_data[c_idx]
                            else:
                                val = ""
                            
                            # Standardize line breaks inside table cells
                            val = val.replace('<br>', '\n')
                            
                            cell = table.cell(r_idx, c_idx)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.line_spacing = 1.15
                            
                            # Format title row (bold, colored background)
                            if r_idx == 0:
                                run = p.add_run(val)
                                run.bold = True
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                                set_cell_background(cell, "ECECEC") # light grey
                            else:
                                parse_inline(p, val)
                                for run in p.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(11)
                                
                    # Add an empty paragraph after the table for spacing
                    doc.add_paragraph()
                
                table_lines = []
            continue

        # 3. Handle Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(title_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(20)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif level == 2:
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(title_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif level == 3:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(title_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(title_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # 4. Handle Blockquotes (Alerts / Figures)
        if line.startswith('>'):
            quote_text = line.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            parse_inline(p, quote_text)
            set_font_and_spacing(p, space_after_pt=6)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # 5. Handle Horizontal Rules
        if line == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # 6. Handle Bullet Lists
        if line.startswith('- ') or line.startswith('* '):
            list_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run("•  ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            parse_inline(p, list_text)
            for r in p.runs[1:]:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
            i += 1
            continue

        # 7. Handle Regular Paragraphs
        if line:
            # Check if this line is part of a numbered list like "1. Item"
            if re.match(r'^\d+\.\s', line):
                add_paragraph_with_font(doc, line)
            else:
                add_paragraph_with_font(doc, line)
        else:
            # Empty line
            pass
            
        i += 1

    try:
        doc.save(output_path)
        print(f"Successfully generated {output_path}")
    except PermissionError:
        alternative_path = os.path.join(workspace_dir, "v2x_project_report_final.docx")
        doc.save(alternative_path)
        print(f"Successfully generated alternative: {alternative_path}")

if __name__ == "__main__":
    main()
