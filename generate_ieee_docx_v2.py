import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def parse_inline(paragraph, text, font_name="Times New Roman", font_size_pt=10, italic_all=False):
    """Parse inline markdown: bold (**text**), italic (*text*), math ($math$), links ([text](url)), code (`code`)"""
    pattern = re.compile(r'(\*\*.*?\*\*|\*[^*]+?\*|\$[^$]+?\$|\[.*?\]\(.*?\)|`[^`]+?`)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
            if italic_all:
                run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(font_size_pt - 1)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
        elif part.startswith('$') and part.endswith('$'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.name = 'Cambria Math'
            run.font.size = Pt(font_size_pt)
        elif part.startswith('[') and ']' in part and part.endswith(')'):
            match = re.search(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text = match.group(1)
                run = paragraph.add_run(link_text)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 238)
                run.font.name = font_name
                run.font.size = Pt(font_size_pt)
                if italic_all:
                    run.font.italic = True
        else:
            run = paragraph.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
            if italic_all:
                run.font.italic = True

def main():
    workspace_dir = r"c:\Users\AYUSH K BHAT\OneDrive\Desktop\v2v-v2i-project-phase2"
    
    md_path = os.path.join(workspace_dir, "syntrix_v2x_ieee_paper_v2.md")
    output_path = os.path.join(workspace_dir, "syntrix_v2x_ieee_paper_v2.docx")

    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = docx.Document()

    # ── IEEE Page Setup ──
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    lines = md_content.split('\n')
    
    # ── Extract Title ──
    title_raw = ""
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title_raw = stripped.lstrip('#').strip()
            break

    # ── Extract Abstract ──
    abstract_raw = ""
    keywords_raw = ""
    in_abstract = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('## Abstract'):
            in_abstract = True
            continue
        if in_abstract:
            if stripped.startswith('**Keywords'):
                kw_match = re.search(r'\*\*Keywords.*?[—\-](.*?)\.\*\*', stripped)
                if kw_match:
                    keywords_raw = kw_match.group(1).strip()
                else:
                    keywords_raw = stripped.replace('**Keywords—', '').replace('**Keywords -', '').replace('**', '').strip()
                in_abstract = False
                continue
            if stripped.startswith('---') or stripped.startswith('## '):
                in_abstract = False
                continue
            if stripped:
                abstract_raw += stripped + " "

    # ── Extract Authors ──
    authors_raw = ""
    in_authors = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('## Authors'):
            in_authors = True
            continue
        if in_authors:
            if stripped.startswith('---') or stripped.startswith('## '):
                break
            if stripped and not stripped.startswith('>'):
                authors_raw += stripped.replace('**', '') + "\n"

    authors_raw = authors_raw.strip() or "Ayush K. Bhat, Vishal\nDepartment of Computer Science and Engineering\nBengaluru, Karnataka, India"

    # ── Build Document ──

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run(title_raw)
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0, 0, 0)

    # Authors
    for author_line in authors_raw.split('\n'):
        author_line = author_line.strip()
        if not author_line:
            continue
        p_auth = doc.add_paragraph()
        p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_auth.paragraph_format.space_after = Pt(2)
        
        is_italic = author_line.startswith('*') and author_line.endswith('*')
        clean_text = author_line.strip('*')
        
        run = p_auth.add_run(clean_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if is_italic:
            run.font.italic = True
    
    # Spacer after authors
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(8)

    # Abstract
    p_abs = doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_abs.paragraph_format.left_indent = Inches(0.25)
    p_abs.paragraph_format.right_indent = Inches(0.25)
    p_abs.paragraph_format.space_after = Pt(6)
    
    run_abs_tag = p_abs.add_run("Abstract\u2014")
    run_abs_tag.bold = True
    run_abs_tag.font.italic = True
    run_abs_tag.font.name = 'Times New Roman'
    run_abs_tag.font.size = Pt(9)
    
    parse_inline(p_abs, abstract_raw.strip(), font_name="Times New Roman", font_size_pt=9, italic_all=True)

    # Keywords
    if keywords_raw:
        p_key = doc.add_paragraph()
        p_key.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_key.paragraph_format.left_indent = Inches(0.25)
        p_key.paragraph_format.right_indent = Inches(0.25)
        p_key.paragraph_format.space_after = Pt(18)
        
        run_key_tag = p_key.add_run("Keywords\u2014")
        run_key_tag.bold = True
        run_key_tag.font.italic = True
        run_key_tag.font.name = 'Times New Roman'
        run_key_tag.font.size = Pt(9)
        
        run_key_val = p_key.add_run(keywords_raw)
        run_key_val.font.name = 'Times New Roman'
        run_key_val.font.size = Pt(9)

    # ── Continuous Section Break → 2 Columns ──
    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    body_section.top_margin = Inches(0.75)
    body_section.bottom_margin = Inches(0.75)
    body_section.left_margin = Inches(0.75)
    body_section.right_margin = Inches(0.75)
    
    sectPr = body_section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '360')
    sectPr.append(cols)

    # ── Process Body Lines ──
    i = 0
    in_code_block = False
    code_text = ""
    table_lines = []
    
    # Skip metadata sections we already processed
    skip_sections = {'Abstract', 'Authors'}
    current_skip = None
    
    while i < len(lines):
        line = lines[i].strip()
        raw_line = lines[i]
        
        # Skip title line
        if line.startswith('# ') and not line.startswith('## '):
            i += 1
            continue
        
        # Skip blockquote metadata lines
        if line.startswith('> **IEEE Paper Format'):
            i += 1
            while i < len(lines) and lines[i].strip().startswith('>'):
                i += 1
            continue
        if line.startswith('>') and ('IEEEtran' in line or 'Following' in line):
            i += 1
            continue
        
        # Skip sections we already processed
        if line.startswith('## Abstract') or line.startswith('## Authors'):
            current_skip = line.replace('## ', '')
            i += 1
            continue
        if current_skip:
            if line.startswith('## ') and not line.startswith('## Abstract') and not line.startswith('## Authors'):
                current_skip = None
                # Fall through to process this heading
            elif line.startswith('---'):
                current_skip = None
                i += 1
                continue
            else:
                i += 1
                continue
        
        # Skip keyword line (already handled)
        if line.startswith('**Keywords'):
            i += 1
            continue
        
        # ── Code Blocks ──
        if line.startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.1)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(code_text.strip())
                run.font.name = 'Consolas'
                run.font.size = Pt(7.5)
                in_code_block = False
                code_text = ""
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_text += lines[i] + "\n"
            i += 1
            continue

        # ── Tables ──
        if line.startswith('|'):
            table_lines.append(lines[i])
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            if table_lines:
                parsed_rows = []
                for t_line in table_lines:
                    content = t_line.strip()
                    if content.startswith('|'):
                        content = content[1:]
                    if content.endswith('|'):
                        content = content[:-1]
                    cells = [c.strip() for c in content.split('|')]
                    is_sep = all(re.match(r'^[\s:-]*$', c) for c in cells)
                    if not is_sep:
                        parsed_rows.append(cells)
                
                if parsed_rows:
                    num_rows = len(parsed_rows)
                    num_cols = max(len(row) for row in parsed_rows)
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.style = 'Table Grid'
                    
                    for r_idx, row_data in enumerate(parsed_rows):
                        for c_idx in range(num_cols):
                            val = row_data[c_idx] if c_idx < len(row_data) else ""
                            val = val.replace('<br>', '\n')
                            
                            cell = table.cell(r_idx, c_idx)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.line_spacing = 1.0
                            
                            if r_idx == 0:
                                run = p.add_run(val.replace('**', ''))
                                run.bold = True
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(8)
                                set_cell_background(cell, "F2F2F2")
                            else:
                                parse_inline(p, val, font_size_pt=8)
                    doc.add_paragraph().paragraph_format.space_after = Pt(4)
                table_lines = []
            continue

        # ── Image references (![caption](path)) ──
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', line)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            
            # Add image if file exists
            if os.path.exists(img_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(3.2))
                except Exception:
                    pass
            
            # Add caption
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(8)
            run_cap = p_cap.add_run(caption)
            run_cap.font.italic = True
            run_cap.font.name = 'Times New Roman'
            run_cap.font.size = Pt(8)
            i += 1
            continue

        # ── Headings ──
        if line.startswith('## '):
            heading_text = line.lstrip('#').strip()
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            # Primary section (Roman numeral) → Centered, Bold, CAPS
            if re.match(r'^[IXV]+\.\s', heading_text) or heading_text in ('References', 'Acknowledgments', 'Author Biographies'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(heading_text.upper())
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(heading_text)
                run.font.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue
            
        if line.startswith('### '):
            heading_text = line.lstrip('#').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            
            run = p.add_run(heading_text)
            run.font.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # ── Blockquotes ──
        if line.startswith('>'):
            quote_text = line.lstrip('>').strip()
            if not quote_text:
                i += 1
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parse_inline(p, quote_text, font_size_pt=9)
            for run in p.runs:
                run.font.italic = True
            i += 1
            continue

        # ── Numbered Lists ──
        numbered_match = re.match(r'^(\d+)\.\s+(.*)', line)
        if numbered_match and not line.startswith('['):
            num = numbered_match.group(1)
            list_text = numbered_match.group(2)
            
            # Check if this is a reference (in References section)
            # References have [N] format in our paper
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parse_inline(p, f"{num}. {list_text}", font_size_pt=10)
            i += 1
            continue

        # ── References [N] format ──
        ref_match = re.match(r'^\[(\d+)\]\s+(.*)', line)
        if ref_match:
            ref_text = line
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parse_inline(p, ref_text, font_size_pt=8)
            i += 1
            continue

        # ── Bullet Lists ──
        if line.startswith('- ') or line.startswith('* '):
            list_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run("\u2022  ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            parse_inline(p, list_text, font_size_pt=10)
            i += 1
            continue
        
        # ── Sub-bullets ──
        if re.match(r'^\s+[-*]\s', raw_line):
            list_text = raw_line.strip().lstrip('-*').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run("\u2013  ")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            parse_inline(p, list_text, font_size_pt=9)
            i += 1
            continue

        # ── Horizontal Rules ──
        if line == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # ── Equations $$ eq $$ ──
        if line.startswith('$$') and line.endswith('$$'):
            eq_text = line[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(eq_text)
            run.font.name = 'Cambria Math'
            run.font.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # ── Standalone equations with numbering ──
        eq_num_match = re.match(r'^\$\$(.*?)\$\$\s*$', line)
        if eq_num_match:
            eq_text = eq_num_match.group(1).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(eq_text)
            run.font.name = 'Cambria Math'
            run.font.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # ── Normal Paragraphs ──
        if line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.0
            parse_inline(p, line, font_size_pt=10.0)
            
        i += 1

    # ── Save ──
    try:
        doc.save(output_path)
        print(f"✅ Successfully generated IEEE paper DOCX at: {output_path}")
    except PermissionError:
        alternative_path = os.path.join(workspace_dir, "syntrix_v2x_ieee_paper_v2_final.docx")
        doc.save(alternative_path)
        print(f"✅ Generated alternative: {alternative_path}")

    print(f"\n📊 Paper Statistics:")
    word_count = len(md_content.split())
    print(f"   Words: ~{word_count}")
    print(f"   Pages: ~{word_count // 700} (estimated for 2-column IEEE)")

if __name__ == "__main__":
    main()
